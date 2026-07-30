from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE
from pathlib import Path

ROOT = Path(r"C:\Users\Joao\Documents\siteeee2")
OUT = ROOT / "Documentacao_FozGames.docx"
IMAGES = [Path(rf"C:\Users\Joao\Downloads\{n}.png") for n in ["6", "1", "2", "5", "3", "4"]]

NAVY = "202A44"
ORANGE = "EF613D"
INK = "171717"
MUTED = "666666"
PALE = "F3F0E8"

def set_font(run, name="Aptos", size=None, color=None, bold=None, italic=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    if size: run.font.size = Pt(size)
    if color: run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None: run.bold = bold
    if italic is not None: run.italic = italic

def shade(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), fill)
    tcPr.append(shd)

def set_cell_margins(cell, top=100, start=120, bottom=100, end=120):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.first_child_found_in("w:tcMar")
    if tcMar is None:
        tcMar = OxmlElement('w:tcMar'); tcPr.append(tcMar)
    for m, v in (("top", top),("start", start),("bottom", bottom),("end", end)):
        node = tcMar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}"); tcMar.append(node)
        node.set(qn('w:w'), str(v)); node.set(qn('w:type'), 'dxa')

def set_repeat_table_header(row):
    trPr = row._tr.get_or_add_trPr(); el = OxmlElement('w:tblHeader'); el.set(qn('w:val'), 'true'); trPr.append(el)

def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = paragraph.add_run('FOZGAMES  |  '); set_font(r, size=8, color=MUTED, bold=True)
    fld = OxmlElement('w:fldSimple'); fld.set(qn('w:instr'), 'PAGE')
    paragraph._p.append(fld)

doc = Document()
sec = doc.sections[0]
sec.top_margin = Inches(0.8); sec.bottom_margin = Inches(0.75)
sec.left_margin = Inches(0.85); sec.right_margin = Inches(0.85)
sec.header_distance = Inches(0.35); sec.footer_distance = Inches(0.35)

styles = doc.styles
normal = styles['Normal']; normal.font.name = 'Aptos'; normal._element.rPr.rFonts.set(qn('w:ascii'), 'Aptos'); normal.font.size = Pt(10.5); normal.font.color.rgb = RGBColor.from_string(INK); normal.paragraph_format.space_after = Pt(6); normal.paragraph_format.line_spacing = 1.15
for sty, size, color, before, after in [('Heading 1',18,NAVY,16,7),('Heading 2',13,ORANGE,12,5),('Heading 3',11,NAVY,9,4)]:
    s=styles[sty]; s.font.name='Aptos Display'; s._element.rPr.rFonts.set(qn('w:ascii'),'Aptos Display'); s.font.size=Pt(size); s.font.color.rgb=RGBColor.from_string(color); s.font.bold=True; s.paragraph_format.space_before=Pt(before); s.paragraph_format.space_after=Pt(after)

footer = sec.footer.paragraphs[0]; add_page_number(footer)

# Cover: customer-pack inspired.
p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(64); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run('DOCUMENTACAO DO SISTEMA'); set_font(r, size=10, color=ORANGE, bold=True)
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after=Pt(8)
r=p.add_run('FozGames'); set_font(r, name='Aptos Display', size=34, color=NAVY, bold=True)
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after=Pt(42)
r=p.add_run('Loja digital de jogos novos e usados'); set_font(r, size=15, color=MUTED)
table=doc.add_table(rows=3, cols=2); table.alignment=WD_TABLE_ALIGNMENT.CENTER; table.autofit=False
metadata=[('Projeto','FozGames'),('Desenvolvedor','Joao Victor dos Santos'),('Repositorio','github.com/jo-4o/siteeee2'),('Tecnologias','React, Vite, JavaScript, CSS e localStorage'),('Versao','1.0'),('Data','Julho de 2026')]
for i,(label,value) in enumerate(metadata):
    cell=table.cell(i//2,(i%2)); shade(cell, PALE); set_cell_margins(cell)
    p=cell.paragraphs[0]; p.paragraph_format.space_after=Pt(2); r=p.add_run(label.upper()); set_font(r,size=7.5,color=ORANGE,bold=True)
    p=cell.add_paragraph(); p.paragraph_format.space_after=Pt(0); r=p.add_run(value); set_font(r,size=10,color=INK,bold=True)
doc.add_paragraph('')
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run('Este documento apresenta o objetivo, as funcionalidades, a estrutura tecnica e as telas do sistema desenvolvido.'); set_font(r,size=10.5,color=MUTED,italic=True)
doc.add_page_break()

def add_bullets(items):
    for item in items:
        p=doc.add_paragraph(style='List Bullet'); p.paragraph_format.space_after=Pt(3); p.add_run(item)

doc.add_heading('1. Visao geral', level=1)
doc.add_paragraph('O FozGames e uma aplicacao web para uma loja de jogos novos e usados. O sistema oferece uma vitrine de jogos para clientes e uma area administrativa para cadastrar produtos e acompanhar os pedidos reservados.')
doc.add_heading('1.1 Objetivo', level=2)
doc.add_paragraph('Digitalizar a apresentacao dos jogos da loja, facilitar a reserva de pedidos e centralizar o controle basico de estoque em uma interface responsiva e simples de utilizar.')
doc.add_heading('1.2 Publico e perfis de acesso', level=2)
add_bullets(['Cliente: acessa o catalogo, filtra jogos, adiciona produtos ao carrinho e reserva pedidos.', 'Administrador: acessa as funcionalidades do cliente e tambem pode cadastrar jogos, enviar capas, remover itens e consultar pedidos.', 'Controle de acesso: o painel de estoque e a pagina de pedidos so aparecem quando o login foi realizado como administrador.'])
doc.add_heading('1.3 Repositorio e execucao', level=2)
doc.add_paragraph('Repositorio: https://github.com/jo-4o/siteeee2')
doc.add_paragraph('Para executar localmente: npm install e npm run dev. Em seguida, acessar o endereco exibido pelo Vite, normalmente http://localhost:5173.')

doc.add_heading('2. Tecnologias e estrutura de dados', level=1)
doc.add_paragraph('A interface foi implementada em React com JavaScript, utilizando Vite como ambiente de desenvolvimento e CSS para estilizacao. A aplicacao nao depende de banco de dados externo para a demonstracao.')
table=doc.add_table(rows=1, cols=3); table.style='Table Grid'; table.alignment=WD_TABLE_ALIGNMENT.CENTER
hdr=table.rows[0].cells; set_repeat_table_header(table.rows[0])
for c,text in zip(hdr,['Estrutura','Campos principais','Finalidade']):
    shade(c, NAVY); set_cell_margins(c); r=c.paragraphs[0].add_run(text); set_font(r,size=9,color='FFFFFF',bold=True)
rows=[('Jogo','id, nome, plataforma, categoria, condicao, preco, estoque, capa','Representa cada item disponivel no catalogo.'),('Usuario','nome, e-mail, perfil','Mantem a sessao de cliente ou administrador.'),('Carrinho','jogo, quantidade','Armazena itens selecionados antes da reserva.'),('Pedido','numero, cliente, contato, endereco, itens, total, data','Registra a reserva feita pelo cliente.')]
for row in rows:
    cells=table.add_row().cells
    for c,text in zip(cells,row):
        set_cell_margins(c); c.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER; r=c.paragraphs[0].add_run(text); set_font(r,size=9)

doc.add_heading('2.1 Persistencia local', level=2)
doc.add_paragraph('As informacoes temporarias sao persistidas no localStorage do navegador. Sao utilizadas as chaves fg_user, fg_games, fg_cart e fg_orders. Dessa forma, os dados permanecem disponiveis entre recarregamentos no mesmo navegador e dispositivo.')
doc.add_heading('2.2 Regras principais', level=2)
add_bullets(['O cliente precisa informar endereco e telefone para concluir a reserva.', 'O carrinho respeita o estoque disponivel ao aumentar quantidades.', 'A reserva nao realiza pagamento on-line; ela registra os dados para combinacao de retirada ou entrega.', 'Imagens de capa podem ser enviadas no cadastro do produto nos formatos JPG, PNG ou WEBP, limitadas a 1,5 MB.', 'Pedidos so sao mostrados na pagina administrativa depois que um cliente conclui a reserva.'])

doc.add_page_break()
doc.add_heading('3. Funcionalidades do sistema', level=1)
features=[('Login e perfis','A pagina inicial permite entrar como cliente ou administrador. O login de administrador libera o menu de Estoque e Pedidos.'),('Catalogo','Exibe os jogos cadastrados em cartoes e permite filtrar por plataforma e genero.'),('Carrinho e reserva','Permite adicionar ou remover itens, ajustar quantidades, informar contato e reservar o pedido.'),('Estoque','Permite ao administrador adicionar jogos, definir plataforma, categoria, condicao, preco, quantidade e enviar uma capa.'),('Pedidos','Apresenta quantidade de pedidos, valor reservado e os dados das reservas feitas pelos clientes.')]
for name,desc in features:
    doc.add_heading(name,level=2); doc.add_paragraph(desc)
doc.add_heading('4. Evidencias das telas', level=1)
doc.add_paragraph('As figuras a seguir registram as principais telas e fluxos implementados no sistema.')

captions=[
 'Figura 1 - Tela de login, com escolha entre os perfis Cliente e Administrador.',
 'Figura 2 - Catalogo do cliente com filtros por plataforma e genero.',
 'Figura 3 - Carrinho vazio, apresentando a acao para retornar a loja.',
 'Figura 4 - Carrinho com produto, total do pedido e formulario de reserva.',
 'Figura 5 - Painel administrativo de estoque com cadastro e envio de capa do jogo.',
 'Figura 6 - Area administrativa de pedidos com indicadores e reservas registradas.'
]
for idx,(img,caption) in enumerate(zip(IMAGES,captions),1):
    if idx > 1: doc.add_page_break()
    doc.add_heading(f'4.{idx} {caption.split(" - ",1)[1]}',level=2)
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(img), width=Inches(6.4))
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after=Pt(0)
    r=p.add_run(caption); set_font(r,size=8.5,color=MUTED,italic=True)

doc.add_page_break()
doc.add_heading('5. Validacao e proximos passos', level=1)
doc.add_paragraph('A aplicacao atende ao fluxo definido para a demonstracao: autenticar um usuario, explorar jogos, montar o carrinho, reservar um pedido e consultar a reserva pela area administrativa. O uso do localStorage atende a necessidade de persistencia temporaria no navegador.')
doc.add_heading('5.1 Roteiro sugerido para demonstracao', level=2)
add_bullets(['Entrar como cliente, informar nome e e-mail.', 'Explorar o catalogo e aplicar um filtro.', 'Adicionar um jogo ao carrinho e concluir uma reserva com endereco e telefone.', 'Sair e entrar como administrador.', 'Abrir Pedidos para confirmar o pedido reservado.', 'Abrir Estoque e demonstrar o cadastro de um novo jogo com capa.'])
doc.add_heading('5.2 Limitacoes atuais', level=2)
add_bullets(['Os dados ficam restritos ao navegador e dispositivo em que foram registrados.', 'O pagamento on-line nao faz parte desta versao.', 'O login administrativo usa credenciais demonstrativas, adequadas apenas ao prototipo academico.'])
doc.add_page_break()
doc.add_heading('6. Gestao das atividades e controle de versao', level=1)
doc.add_paragraph('As atividades do projeto foram organizadas em um quadro de gestao, com listas para acompanhar itens a fazer, em andamento e concluídos. As tarefas de login, funcionalidades da loja, estilos, configuracao local, documentacao e publicacao do repositorio foram registradas como concluidas.')
doc.add_heading('6.1 Ferramenta de gestao', level=2)
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
p.add_run().add_picture(r'C:\Users\Joao\AppData\Local\Temp\codex-clipboard-e0f91634-5c40-4f71-9573-f7b48569ebcb.png', width=Inches(6.4))
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run('Figura 7 - Quadro de gestao com as atividades do projeto marcadas como concluidas.'); set_font(r,size=8.5,color=MUTED,italic=True)
doc.add_page_break()
doc.add_heading('6.2 Historico de commits', level=2)
doc.add_paragraph('O codigo-fonte foi versionado em um repositorio GitHub. Os commits foram separados por etapa de desenvolvimento, facilitando o acompanhamento das alteracoes e a comprovacao da evolucao do projeto.')
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
p.add_run().add_picture(r'C:\Users\Joao\AppData\Local\Temp\codex-clipboard-4292c4aa-472c-4ece-b108-26633e32dd4e.png', width=Inches(4.8))
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run('Figura 8 - Historico de commits no repositorio GitHub do projeto.'); set_font(r,size=8.5,color=MUTED,italic=True)
doc.add_paragraph('Fim da documentacao.')

doc.core_properties.title = 'Documentacao do Site - FozGames'
doc.core_properties.author = 'Joao Victor dos Santos'
doc.save(OUT)
print(OUT)
